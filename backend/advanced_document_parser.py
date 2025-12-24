import os
import io
import shutil
import logging
import tempfile
import subprocess
import uuid
import time
import requests
from pathlib import Path
from urllib.parse import urlparse

from config import config



# === 依赖库检查与导入 ===
try:
    import pytesseract
except ImportError:
    print("[ERROR] 警告: 未安装 pytesseract")

try:
    # 关键优化：使用 convert_from_path 而不是 convert_from_bytes，节省内存
    from pdf2image import convert_from_path
except ImportError:
    print("[ERROR] 警告: 未安装 pdf2image")

try:
    from pypdf import PdfReader
    HAS_PDF = True
except ImportError:
    print("[ERROR] 警告: 未安装 pypdf")
    HAS_PDF = False

try:
    import docx
    from docx import Document
except ImportError:
    print("[ERROR] 警告: 未安装 python-docx")

try:
    # [新增] Excel 支持
    import pandas as pd
    import openpyxl  # pandas 依赖

    HAS_EXCEL = True
except ImportError:
    HAS_EXCEL = False
    print("[WARNING] 提示: 未安装 pandas/openpyxl，无法解析 .xlsx")


class AdvancedDocumentProcessor:
    """
    终极文档解析器 (v2)
    集成：LibreOffice (.doc支持), Pandas (Excel支持), OCR (扫描件支持), 流式下载 (内存安全)
    """

    def __init__(self, ocr_lang='chi_sim+eng', max_ocr_pages=5, temp_dir=None):
        self.ocr_lang = ocr_lang
        self.max_ocr_pages = max_ocr_pages
        # 确保临时目录存在
        self.temp_dir = temp_dir if temp_dir else tempfile.gettempdir()

        self.supported_file_types = config.SUPPORTED_FILE_TYPES

        logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(name)s - %(levelname)s - %(message)s')
        self.logger = logging.getLogger("DocParser")
        self.logger.setLevel(logging.INFO)

        # 自动检测 LibreOffice 路径
        self.libreoffice_path = self._find_libreoffice()
        if not self.libreoffice_path:
            self.logger.warning("⚠️ 未检测到 LibreOffice，.doc 文件解析能力将受限")
        else:
            self.logger.info(f"✅ LibreOffice 就绪: {self.libreoffice_path}")

    def extract_content(self, file_url_or_path, file_type=None, file_name=""):
        """
        解析主入口
        """
        local_path = None
        is_temp = False

        try:
            # 1. 统一获取本地文件路径 (URL则下载，流则存盘)
            # 生产环境原则：尽量操作磁盘文件，避免操作巨大内存流
            local_path, is_temp = self._resolve_file_path(file_url_or_path)
            if not local_path or not os.path.exists(local_path):
                return "❌ 文件获取失败或不存在"

            # 2. 智能推断类型
            ext = file_type if file_type else os.path.splitext(local_path)[-1].replace('.', '').lower()
            if not ext and file_name:
                ext = os.path.splitext(file_name)[-1].replace('.', '').lower()

            ext = ext.lower()

            # 3. 分发处理
            if ext not in ["pdf","docx",".pdf",".docx"]:
                return "暂不支持的文件格式"

            if 'pdf' in ext:
                return self._parse_pdf(local_path)
            elif ext in ['jpg', 'jpeg', 'png', 'bmp']:
                return self._parse_image(local_path)
            elif 'docx' == ext:
                return self._parse_docx(local_path)
            elif 'doc' == ext:
                return self._parse_doc_legacy(local_path)  # 调用 LibreOffice
            elif ext in ['xlsx', 'xls', 'csv']:
                return self._parse_excel(local_path, ext)
            elif 'txt' in ext or 'md' in ext:
                with open(local_path, 'r', encoding='utf-8', errors='ignore') as f:
                    return f.read()
            else:
                return f"暂不支持的文件格式: {ext}"

        except Exception as e:
            self.logger.error(f"解析异常: {e}", exc_info=True)
            return f"解析发生系统错误: {str(e)}"
        finally:
            # 4. 清理临时文件
            if is_temp and local_path and os.path.exists(local_path):
                try:
                    os.remove(local_path)
                    self.logger.info(f"清理临时文件: {local_path}")
                except Exception as e:
                    self.logger.warning(f"清理临时文件失败: {e}")


    def _resolve_file_path(self, path_or_url):
        """
        获取可操作的本地文件路径。
        如果是 URL，下载到临时文件。
        如果是 本地路径，直接返回。
        返回: (file_path, is_temporary_file)
        """
        path_str = str(path_or_url).strip()

        # 检查是否是 URL
        if path_str.startswith(('http://', 'https://')):
            # [安全] SSRF 简单防护：禁止访问内网 IP (此处仅为示例，生产环境需更严格的校验)
            domain = urlparse(path_str).netloc
            if 'localhost' in domain or '127.0.0.1' in domain or '192.168.' in domain:
                raise ValueError("禁止访问内网地址")

            # 下载流式写入，防止内存溢出
            try:
                with requests.get(path_str, stream=True, timeout=30) as r:
                    r.raise_for_status()
                    # 创建临时文件
                    suffix = os.path.splitext(urlparse(path_str).path)[-1]
                    if not suffix: suffix = ".tmp"

                    tf = tempfile.NamedTemporaryFile(delete=False, dir=self.temp_dir, suffix=suffix)
                    for chunk in r.iter_content(chunk_size=8192):
                        tf.write(chunk)
                    tf.close()
                    return tf.name, True
            except Exception as e:
                self.logger.error(f"URL 下载失败: {e}")
                return None, False

        # 本地文件
        if path_str.lower().startswith('file://'):
            path_str = path_str[7:]

        if os.path.exists(path_str):
            return path_str, False

        return None, False


    # ==========================================
    # 核心解析逻辑
    # ==========================================

    def _parse_doc_legacy(self, file_path):
        """
        解析 .doc 文件 (使用 LibreOffice 转 docx)
        修复：使用 Path.as_uri() 解决 Windows 下 bootstrap.ini 报错问题
        """
        if not self.libreoffice_path:
            return "⚠️ 服务器未安装 LibreOffice，无法解析 .doc 格式。请另存为 .docx 或 .pdf 上传。"

        # 创建临时目录
        temp_dir = tempfile.mkdtemp()
        try:
            self.logger.info("启动 LibreOffice 转换 .doc -> .docx ...")

            # [修复关键点]
            # Windows 下必须使用 file:///C:/... 格式，且需要 URL 编码
            # Path(...).as_uri() 会自动处理 Windows 路径格式和转义
            user_profile_dir = os.path.join(temp_dir, "user_profile")
            user_profile_url = Path(user_profile_dir).as_uri()

            cmd = [
                self.libreoffice_path,
                "--headless",
                "--convert-to", "docx",
                "--outdir", temp_dir,
                f"-env:UserInstallation={user_profile_url}",  # 使用标准 URI
                file_path
            ]

            # Windows 下防止弹出 CMD 窗口 (可选优化)
            startupinfo = None
            if os.name == 'nt':
                startupinfo = subprocess.STARTUPINFO()
                startupinfo.dwFlags |= subprocess.STARTF_USESHOWWINDOW

            subprocess.run(
                cmd,
                check=True,
                timeout=60,
                stdout=subprocess.DEVNULL,
                stderr=subprocess.DEVNULL,
                startupinfo=startupinfo
            )

            # 2. 找结果文件
            original_name = Path(file_path).stem
            converted_path = os.path.join(temp_dir, f"{original_name}.docx")

            if os.path.exists(converted_path):
                return self._parse_docx(converted_path)
            else:
                # 模糊匹配防止文件名乱码或被截断
                files = os.listdir(temp_dir)
                docx_files = [f for f in files if f.endswith('.docx')]
                if docx_files:
                    return self._parse_docx(os.path.join(temp_dir, docx_files[0]))
                return "❌ .doc 转换失败，未生成目标文件"

        except subprocess.TimeoutExpired:
            return "❌ .doc 解析超时 (文件可能过大或损坏)"
        except Exception as e:
            return f"❌ .doc 转换出错: {e}"
        finally:
            # 清理临时目录
            try:
                shutil.rmtree(temp_dir, ignore_errors=True)
            except Exception as e:
                self.logger.warning(f"清理临时目录失败: {e}")

    def _parse_docx(self, file_path):
        """解析 Word"""
        doc = docx.Document(file_path)
        # 优化：同时提取表格内容
        text_parts = []

        # 1. 段落
        for p in doc.paragraphs:
            if p.text.strip():
                text_parts.append(p.text)

        # 2. 表格 (工程文档重要信息常在表格)
        if doc.tables:
            text_parts.append("\n--- 表格内容 ---\n")
            for table in doc.tables:
                for row in table.rows:
                    row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                    if row_text:
                        text_parts.append(" | ".join(row_text))

        return "\n".join(text_parts)

    def _parse_excel(self, file_path, ext):
        """[新增] 解析 Excel/CSV"""
        if not HAS_EXCEL:
            return "❌ 服务器缺少 pandas 库，无法解析 Excel"

        try:
            content = []
            if 'csv' in ext:
                df = pd.read_csv(file_path)
                content.append(df.to_string(index=False))
            else:
                # 读取所有 sheet
                dfs = pd.read_excel(file_path, sheet_name=None)
                for sheet_name, df in dfs.items():
                    content.append(f"\n=== 工作表: {sheet_name} ===\n")
                    # 转换为 markdown 风格或纯文本表格
                    # fillna('') 防止 None 显示为 NaN
                    content.append(df.fillna('').to_string(index=False))

            return "\n".join(content)
        except Exception as e:
            return f"❌ Excel 解析失败: {e}"

    def _parse_image(self, file_path):
        """解析图片 (OCR)"""
        try:
            from PIL import Image
            img = Image.open(file_path)
            # 限制图片大小防止 OCR 耗时过长
            if img.width * img.height > 20000000:  # > 20MP
                self.logger.warning("图片过大，进行缩放...")
                img.thumbnail((2048, 2048))

            text = pytesseract.image_to_string(img, lang=self.ocr_lang)
            return f"【图片OCR】\n{text}"
        except Exception as e:
            return f"❌ 图片OCR失败: {e}"

    def _parse_pdf(self, file_path):
        """
        PDF 混合解析 (优化版)
        """
        text_content = ""

        # 1. 尝试提取文本
        try:
            if HAS_PDF:
                reader = PdfReader(file_path)
                # 限制页数？通常文本提取很快，不需要限制
                for page in reader.pages:
                    extracted = page.extract_text()
                    if extracted:
                        text_content += extracted + "\n"
            else:
                self.logger.warning("PDF 文本提取失败: 未安装 pypdf 库")
        except Exception as e:
            self.logger.warning(f"PDF 文本提取微弱错误: {e}")

        # 2. 智能判断是否需要 OCR
        # 逻辑：总字数 < 50 且 (有页数 > 0) -> 认为是扫描件
        if len(text_content.strip()) < 50:
            self.logger.info("PDF 为扫描件...")
            ocr_text = ""
            # try:
            #     # 使用 convert_from_path，比 bytes 更省内存
            #     # fmt='jpeg' 比默认 png 更小
            #     from pdf2image import convert_from_path
            #
            #     # 分批处理：避免一次性把几百页全转成图片
            #     # 这里为了简单，我们只取前 N 页，或者你需要写个循环分页读取
            #     images = convert_from_path(
            #         file_path,
            #         first_page=1,
            #         last_page=self.max_ocr_pages,
            #         fmt='jpeg'
            #     )
            #
            #     for i, img in enumerate(images):
            #         self.logger.info(f"OCR 识别第 {i + 1} 页...")
            #         ocr_text += pytesseract.image_to_string(img, lang=self.ocr_lang) + "\n"
            #
            #     if len(images) == self.max_ocr_pages:
            #         ocr_text += f"\n... (为保障性能，仅识别前 {self.max_ocr_pages} 页扫描件) ..."
            #
            #     return f"【PDF扫描件OCR结果】\n{ocr_text}"
            #
            # except Exception as e:
            #     err_msg = str(e)
            #     if "poppler" in err_msg.lower():
            #         return f"❌ OCR 失败: 服务器未安装 Poppler (PDF转图片工具)。\n已提取文本: {text_content}"
            #     return f"❌ OCR 严重错误: {err_msg}"

        return text_content


    # ==========================================
    # 工具函数
    # ==========================================

    def _find_libreoffice(self):
        """跨平台查找 LibreOffice"""
        if os.name == 'nt':
            paths = [
                r"C:\Program Files\LibreOffice\program\soffice.exe",
                r"C:\Program Files (x86)\LibreOffice\program\soffice.exe"
            ]
            for p in paths:
                if os.path.exists(p): return p
            # 尝试 where 命令
            try:
                res = subprocess.run(['where', 'soffice'], capture_output=True, text=True)
                if res.returncode == 0: return res.stdout.split('\n')[0].strip()
            except:
                pass
        else:
            # Linux / Mac
            for cmd in ['libreoffice', 'soffice']:
                try:
                    res = subprocess.run(['which', cmd], capture_output=True, text=True)
                    if res.returncode == 0: return res.stdout.strip()
                except:
                    pass
        return None


# if __name__ == "__main__":
#     p = AdvancedDocumentProcessor()
#     # 测试：请确保有一个 .doc 文件路径
#     # print(p.extract_content("test.doc"))